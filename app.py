import streamlit as st
from pypdf import PdfReader
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import google.generativeai as genai
from google.api_core.exceptions import ResourceExhausted
import io
import time
from datetime import datetime
import re

# ==========================================
# --- 1. BASE DE DADOS JURÍDICA ---
# ==========================================
LEGISLATION_DB = {
    "RJAIA (DL 151-B/2013)": "https://diariodarepublica.pt/dr/legislacao-consolidada/decreto-lei/2013-116043164",
    "Alteração RJAIA (DL 152-B/2017)": "https://diariodarepublica.pt/dr/detalhe/decreto-lei/152-b-2017-114337069",
    "RGGR (DL 102-D/2020)": "https://diariodarepublica.pt/dr/legislacao-consolidada/decreto-lei/2020-150917243",
    "LUA (DL 75/2015)": "https://diariodarepublica.pt/dr/legislacao-consolidada/decreto-lei/2015-106562356",
    "Rede Natura 2000 (DL 140/99)": "https://diariodarepublica.pt/dr/legislacao-consolidada/decreto-lei/1999-34460975",
    "Regulamento Geral do Ruído (DL 9/2007)": "https://diariodarepublica.pt/dr/legislacao-consolidada/decreto-lei/2007-34526556",
    "Lei da Água (Lei 58/2005)": "https://diariodarepublica.pt/dr/legislacao-consolidada/lei/2005-34563267",
    "Emissões Industriais (DL 127/2013)": "https://diariodarepublica.pt/dr/legislacao-consolidada/decreto-lei/2013-34789569"
}

# ==========================================
# --- CONFIGURAÇÃO INICIAL E ESTADO ---
# ==========================================
st.set_page_config(page_title="Análise Caso a Caso RJAIA", page_icon="⚖️", layout="wide")

if 'uploader_key' not in st.session_state:
    st.session_state.uploader_key = 0
if 'validation_result' not in st.session_state:
    st.session_state.validation_result = None
if 'decision_result' not in st.session_state:
    st.session_state.decision_result = None

def reset_app():
    st.session_state.uploader_key += 1
    st.session_state.validation_result = None
    st.session_state.decision_result = None

# ==========================================
# --- SIDEBAR & SETUP ---
# ==========================================
with st.sidebar:
    st.header("🔐 Configuração")
    
    # Gestão de API Key mais robusta
    api_key = None
    if "GOOGLE_API_KEY" in st.secrets:
        api_key = st.secrets["GOOGLE_API_KEY"]
        st.success("Chave API detetada (Secrets)!")
    else:
        api_key = st.text_input("Google API Key", type="password")
    
    selected_model = "gemini-1.5-flash"
    
    if api_key:
        try:
            genai.configure(api_key=api_key)
            models = genai.list_models()
            valid_models = [m.name for m in models if 'generateContent' in m.supported_generation_methods]
            
            if valid_models:
                # Tenta encontrar o flash, se não usa o primeiro disponível
                idx = next((i for i, m in enumerate(valid_models) if 'flash' in m), 0)
                selected_model = st.selectbox("Modelo IA:", valid_models, index=idx)
                st.info("✅ Sistema Pronto")
            else:
                st.warning("⚠️ Chave válida, mas sem modelos disponíveis.")
        except Exception as e:
            st.error(f"Erro na API: {e}")
    
    st.divider()
    if st.button("🔄 Nova Análise / Limpar Tudo", use_container_width=True):
        reset_app()
        st.rerun()

# ==========================================
# --- FUNÇÕES AUXILIARES (WORD) ---
# ==========================================

def add_hyperlink(paragraph, text, url):
    """Adiciona um hiperlink clicável num parágrafo do Word."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), "0000FF")
    rPr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def markdown_to_word(doc, text):
    """Converte Markdown para Word com Justificação."""
    if not text: return
    lines = text.split('\n')
    for line in lines:
        line = line.strip()
        if not line: continue
        
        p = None
        if line.startswith('##'):
            p = doc.add_heading(line.replace('#', '').strip(), level=2)
        elif line.startswith('###'):
            p = doc.add_heading(line.replace('#', '').strip(), level=3)
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            clean_line = line[2:]
            process_bold(p, clean_line)
        else:
            p = doc.add_paragraph()
            process_bold(p, line)
        
        if p and not line.startswith('#'):
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def process_bold(paragraph, text):
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            paragraph.add_run(part[2:-2]).bold = True
        else:
            paragraph.add_run(part)

def append_legislation_section(doc):
    doc.add_page_break()
    doc.add_heading("Legislação Consultada e Referências", level=1)
    
    p_intro = doc.add_paragraph("A presente análise teve por base os seguintes diplomas legais:")
    p_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    for name, url in LEGISLATION_DB.items():
        p = doc.add_paragraph(style='List Bullet')
        add_hyperlink(p, name, url)

# ==========================================
# --- EXTRAÇÃO E IA ---
# ==========================================

def extract_text(files, label):
    """
    Extrai texto de PDFs de forma robusta, convertendo para BytesIO
    para evitar erros de leitura de stream do Streamlit.
    """
    text = ""
    if not files: return "" 
    
    for f in files:
        try:
            # 1. TRUQUE DE SEGURANÇA: Ler o ficheiro para memória primeiro
            # Isso resolve problemas onde o Streamlit entrega o ficheiro "aberto" de forma incorreta
            f.seek(0)
            bytes_data = f.read()
            f_stream = io.BytesIO(bytes_data)
            
            r = PdfReader(f_stream)
            
            # 2. Verificar se está encriptado (comum em D.R.)
            if r.is_encrypted:
                try:
                    r.decrypt("") # Tenta desbloquear se não tiver pass real
                except:
                    st.warning(f"⚠️ O ficheiro '{f.name}' está protegido. A tentar forçar leitura...")

            file_text = ""
            count_pages = len(r.pages)
            
            # 3. Extração página a página
            for i, p in enumerate(r.pages):
                page_content = p.extract_text()
                if page_content:
                    file_text += f"[Pág. {i+1}] {page_content}\n"
            
            # 4. Validação se extraiu algo útil
            if len(file_text.strip()) < 50:
                st.warning(f"⚠️ Atenção: O ficheiro '{f.name}' parece ser uma imagem (digitalização) ou está vazio. A IA não conseguirá lê-lo adequadamente.")
                text += f"\n[AVISO: O ficheiro {f.name} não contém texto selecionável (provável digitalização).]\n"
            else:
                text += f"\n\n>>> FONTE: {label} ({f.name}) <<<\n{file_text}"
                
        except Exception as e:
            st.error(f"❌ Erro crítico ao ler '{f.name}': {str(e)}")
            text += f"\n[ERRO DE LEITURA EM {f.name}: O sistema não conseguiu processar este ficheiro.]\n"
            
    return text

def get_ai(prompt):
    if not api_key:
        return "Erro: Falta a API Key."
        
    model = genai.GenerativeModel(selected_model)
    max_retries = 3
    for attempt in range(max_retries):
        try:
            # Geração com stream=False para garantir resposta completa
            response = model.generate_content(prompt)
            return response.text
        except ResourceExhausted:
            time.sleep(5 * (attempt + 1))
        except Exception as e:
            return f"Erro IA: {str(e)}"
    return "Erro: Sistema sobrecarregado ou erro persistente."

# --- PROMPTS ATUALIZADOS ---

def analyze_validation(t_sim, t_form, t_proj, t_leg):
    legislacao_str = ", ".join(LEGISLATION_DB.keys())
    return get_ai(f"""
    Atua como PERITO AUDITOR AMBIENTAL.
    
    CONTEXTO LEGAL GERAL:
    Utiliza os limiares do RJAIA (Anexos I, II, III, IV, V) e legislação conexa: {legislacao_str}.

    CONTEXTO LEGAL ESPECÍFICO (PDM/Regulamentos):
    {t_leg[:30000]} 

    DADOS DO PROJETO:
    {t_sim[:25000]}
    {t_form[:25000]}
    {t_proj[:80000]}

    TAREFA:
    1. Audita a consistência dos dados (Áreas, Toneladas, LER).
    2. Verifica o enquadramento legal RJAIA.
    3. CRUZAMENTO: Verifica compatibilidade com 'CONTEXTO LEGAL ESPECÍFICO' se existir (ex: índices do PDM, interdições de uso).
    
    OUTPUT (Markdown):
    1. "STATUS: [VALIDADO ou INCONSISTENTE]"
    2. "## 1. Resumo Executivo"
    3. "## 2. Auditoria de Conformidade e Condicionantes Locais"
    4. "## 3. Enquadramento Legal e Limiares"
    """)

def generate_decision_text(t_sim, t_form, t_proj, t_leg):
    return get_ai(f"""
    Atua como Técnico Superior da CCDR. Redige a MINUTA DE DECISÃO.
    
    CONTEXTO DO PROJETO:
    {t_proj[:120000]}
    {t_form[:25000]}
    
    LEGISLAÇÃO ESPECÍFICA / PDM:
    {t_leg[:40000]}

    REGRAS:
    - Texto corrido, JUSTIFICADO, linguagem formal.
    - Cita sempre a fonte: (MD, pág. X).
    
    PREENCHE AS TAGS:
    ### CAMPO_DESIGNACAO
    ### CAMPO_TIPOLOGIA
    ### CAMPO_ENQUADRAMENTO
    ### CAMPO_LOCALIZACAO
    ### CAMPO_AREAS_SENSIVEIS
    ### CAMPO_PROPONENTE
    ### CAMPO_ENTIDADE_LICENCIADORA
    ### CAMPO_AUTORIDADE_AIA
    
    ### CAMPO_DESCRICAO
    (Resumo técnico)
    
    ### CAMPO_CARATERISTICAS
    (Fundamentação técnica quantificada)
    
    ### CAMPO_LOCALIZACAO_PROJETO
    (Compatibilidade com PDM/REN/RAN. Cita especificamente 'LEGISLAÇÃO ESPECÍFICA' se aplicável.)
    
    ### CAMPO_IMPACTES
    (Descritores ambientais)
    
    ### CAMPO_DECISAO
    (SUJEITO / NÃO SUJEITO)
    
    ### CAMPO_CONDICIONANTES
    """)

# ==========================================
# --- GERADORES DE DOCS ---
# ==========================================

def create_validation_doc(text):
    doc = Document()
    
    sec = doc.sections[0]
    sec.header.paragraphs[0].text = "Relatório de Auditoria Técnica"
    sec.header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    h = doc.add_heading("Auditoria de Conformidade Legal e Técnica", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Data: {datetime.now().strftime('%d/%m/%Y')}\n").alignment = WD_ALIGN_PARAGRAPH.CENTER

    if text:
        p_status = doc.add_paragraph()
        p_status.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if "INCONSISTENTE" in text.upper():
            r = p_status.add_run("⚠️ PARECER: INCONGRUÊNCIAS DETETADAS")
            r.font.color.rgb = RGBColor(255, 0, 0)
        else:
            r = p_status.add_run("✅ PARECER: DADOS CONSISTENTES")
            r.font.color.rgb = RGBColor(0, 128, 0)
        r.bold = True
        r.font.size = Pt(14)
        
        doc.add_paragraph("---")
        clean_text = re.sub(r'STATUS:.*', '', text, count=1).strip()
        markdown_to_word(doc, clean_text)
    else:
        doc.add_paragraph("Erro: Sem conteúdo gerado.")
    
    append_legislation_section(doc)
    
    bio = io.BytesIO()
    doc.save(bio)
    return bio

def create_decision_doc(text):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(12) 
    
    if not text:
        doc.add_paragraph("Erro: Sem conteúdo gerado.")
        bio = io.BytesIO()
        doc.save(bio)
        return bio

    def get_tag(tag):
        m = re.search(f"### {tag}(.*?)###", text, re.DOTALL)
        if not m: m = re.search(f"### {tag}(.*)", text, re.DOTALL)
        return m.group(1).strip() if m else "A preencher"

    h = doc.add_heading("Análise prévia e decisão de sujeição a AIA", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'

    def add_merged_header(txt):
        r = table.add_row()
        c = r.cells[0]
        c.merge(r.cells[1])
        p = c.paragraphs[0]
        run = p.add_run(txt)
        run.bold = True
        return r

    def add_row(label, val):
        r = table.add_row()
        p_lbl = r.cells[0].paragraphs[0]
        p_lbl.add_run(label).bold = True
        p_val = r.cells[1].paragraphs[0]
        p_val.text = val
        p_val.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        return r

    def add_full_text(header, content):
        add_merged_header(header)
        r = table.add_row()
        c = r.cells[0]
        c.merge(r.cells[1])
        cell_p = c.paragraphs[0]
        cell_p.clear() 
        paragraphs = content.split('\n')
        first = True
        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text: continue
            if first:
                p = cell_p
                first = False
            else:
                p = c.add_paragraph()
            p.text = para_text
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(12)

    add_merged_header("Identificação")
    add_row("Designação do projeto", get_tag("CAMPO_DESIGNACAO"))
    add_row("Tipologia de Projeto", get_tag("CAMPO_TIPOLOGIA"))
    add_row("Enquadramento no RJAIA", get_tag("CAMPO_ENQUADRAMENTO"))
    add_row("Localização", get_tag("CAMPO_LOCALIZACAO"))
    add_row("Áreas Sensíveis", get_tag("CAMPO_AREAS_SENSIVEIS"))
    add_row("Proponente", get_tag("CAMPO_PROPONENTE"))
    add_row("Entidade Licenciadora", get_tag("CAMPO_ENTIDADE_LICENCIADORA"))
    add_row("Autoridade de AIA", get_tag("CAMPO_AUTORIDADE_AIA"))

    add_full_text("Breve descrição do projeto", get_tag("CAMPO_DESCRICAO"))

    add_merged_header("Fundamentação da decisão")
    add_full_text("Caraterísticas do projeto", get_tag("CAMPO_CARATERISTICAS"))
    add_full_text("Localização do projeto", get_tag("CAMPO_LOCALIZACAO_PROJETO"))
    add_full_text("Impactes Potenciais", get_tag("CAMPO_IMPACTES"))

    add_merged_header("Decisão")
    r = table.add_row()
    c = r.cells[0]
    c.merge(r.cells[1])
    p = c.paragraphs[0]
    run = p.add_run(get_tag("CAMPO_DECISAO"))
    run.bold = True
    run.font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_full_text("Condicionantes", get_tag("CAMPO_CONDICIONANTES"))

    doc.add_paragraph("\n")
    t_sig = doc.add_table(rows=1, cols=2)
    t_sig.rows[0].cells[0].text = "Data: " + datetime.now().strftime('%d/%m/%Y')
    p_sig = t_sig.rows[0].cells[1].paragraphs[0]
    p_sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_sig.add_run("O Técnico,\n\n_______________________").bold = True

    bio = io.BytesIO()
    doc.save(bio)
    return bio

# ==========================================
# --- UI PRINCIPAL ---
# ==========================================
st.title("⚖️ Análise Caso a Caso (RJAIA)")
st.markdown("### Auditoria Técnica e Decisão Fundamentada")

col1, col2, col3, col4 = st.columns(4)

with col1:
    files_sim = st.file_uploader("📂 Simulação SILiAmb", type=['pdf'], accept_multiple_files=True, key=f"s_{st.session_state.uploader_key}")
with col2:
    files_form = st.file_uploader("📂 Formulário", type=['pdf'], accept_multiple_files=True, key=f"f_{st.session_state.uploader_key}")
with col3:
    files_doc = st.file_uploader("📂 Projeto/Memória", type=['pdf'], accept_multiple_files=True, key=f"p_{st.session_state.uploader_key}")
with col4:
    files_leg = st.file_uploader("📜 Legislação/PDM", type=['pdf'], accept_multiple_files=True, key=f"l_{st.session_state.uploader_key}", help="Opcional: PDM, Regulamentos ou Condicionantes Específicas.")

st.markdown("---")

if st.button("🚀 Processar com Rigor Jurídico", type="primary", use_container_width=True):
    # Verificação de ficheiros OBRIGATÓRIOS apenas
    if not (files_sim and files_form and files_doc):
        st.error("⚠️ Atenção: É obrigatório carregar a Simulação, o Formulário e a Memória do Projeto.")
    elif not api_key:
        st.error("🔑 Erro: Chave API em falta na barra lateral.")
    else:
        # Bloco de segurança para apanhar erros (try/except)
        try:
            with st.status("⚙️ A iniciar processamento...", expanded=True) as status:
                
                # 1. Extração de Texto
                st.write("📄 A extrair texto dos documentos...")
                ts = extract_text(files_sim, "SIMULAÇÃO")
                tf = extract_text(files_form, "FORMULÁRIO")
                tp = extract_text(files_doc, "PROJETO")
                
                # Legislação é opcional - usamos ternário para evitar erros se for None
                tl = extract_text(files_leg, "LEGISLAÇÃO_LOCAL") if files_leg else "Nenhuma legislação específica fornecida."
                
                # 2. Validação IA
                st.write("🕵️ A realizar Auditoria e Cruzamento Legal...")
                val_result = analyze_validation(ts, tf, tp, tl)
                
                if "Erro" in val_result and len(val_result) < 100:
                    raise Exception(f"Falha na validação: {val_result}")
                st.session_state.validation_result = val_result
                
                # 3. Decisão IA
                st.write("⚖️ A redigir Minuta de Decisão...")
                dec_result = generate_decision_text(ts, tf, tp, tl)
                
                if "Erro" in dec_result and len(dec_result) < 100:
                    raise Exception(f"Falha na decisão: {dec_result}")
                st.session_state.decision_result = dec_result
                
                status.update(label="✅ Processo Concluído com Sucesso!", state="complete")
        
        except Exception as e:
            st.error(f"❌ Erro Crítico durante a execução: {e}")
            st.info("Sugestão: Verifique se os PDFs não estão corrompidos ou tente novamente.")

# --- ÁREA DE DOWNLOAD ---
if st.session_state.validation_result and st.session_state.decision_result:
    st.success("Documentos gerados com sucesso.")
    c1, c2 = st.columns(2)
    
    f_val = create_validation_doc(st.session_state.validation_result)
    c1.download_button(
        "📄 Baixar Relatório Auditoria (.docx)", 
        f_val.getvalue(), 
        "Relatorio_Auditoria.docx", 
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    
    f_dec = create_decision_doc(st.session_state.decision_result)
    c2.download_button(
        "📝 Baixar Minuta Decisão (.docx)", 
        f_dec.getvalue(), 
        "Proposta_Decisao.docx", 
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
        type="primary"
    )

